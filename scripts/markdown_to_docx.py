from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^-\s+(.*)$")
INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True

    doc.styles["Heading 1"].font.size = Pt(20)
    doc.styles["Heading 2"].font.size = Pt(16)
    doc.styles["Heading 3"].font.size = Pt(13)
    doc.styles["Heading 4"].font.size = Pt(11.5)

    if "Code Inline" not in doc.styles:
        code_style = doc.styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9.5)


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.style = "Code Inline"
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def flush_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return

    rows: list[list[str]] = []
    for raw in table_lines:
        stripped = raw.strip()
        if not stripped:
            continue
        if re.match(r"^\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?$", stripped):
            continue
        if "|" not in stripped:
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for r_index, row in enumerate(rows):
        for c_index in range(col_count):
            value = row[c_index] if c_index < len(row) else ""
            cell = table.cell(r_index, c_index)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, value)
            if r_index == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "EAF2FF")

    doc.add_paragraph("")


def convert_markdown(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    apply_base_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    lines = markdown_text.splitlines()
    table_buffer: list[str] = []

    def flush_pending_table() -> None:
        nonlocal table_buffer
        flush_table(doc, table_buffer)
        table_buffer = []

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()

        if "|" in raw and raw.count("|") >= 2:
            table_buffer.append(raw)
            continue

        if table_buffer:
            flush_pending_table()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped == "---":
            doc.add_paragraph("")
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            hashes, text = heading_match.groups()
            level = min(len(hashes), 4)
            para = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(para, text)
            continue

        ordered_match = ORDERED_RE.match(stripped)
        if ordered_match:
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, ordered_match.group(1))
            continue

        unordered_match = UNORDERED_RE.match(stripped)
        if unordered_match:
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, unordered_match.group(1))
            continue

        if stripped.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            add_inline_runs(para, stripped[2:])
            continue

        para = doc.add_paragraph()
        add_inline_runs(para, stripped)

    if table_buffer:
        flush_pending_table()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python scripts/markdown_to_docx.py <input.md> <output.docx>")
        return 1

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    markdown_text = input_path.read_text(encoding="utf-8")
    convert_markdown(markdown_text, output_path)
    print(f"Wrote DOCX to {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
